from __future__ import annotations

from pathlib import Path

from docx import Document  # type: ignore

from .base import DocumentParser, ParsedDocument


class DocxParser(DocumentParser):
	def parse(self, file_path: str | Path) -> ParsedDocument:
		path = Path(file_path)
		doc = Document(str(path))
		paragraphs = [p.text for p in doc.paragraphs]
		# Include table text as well
		for table in doc.tables:
			for row in table.rows:
				for cell in row.cells:
					paragraphs.append(cell.text)
		raw_text = "\n".join(paragraphs)
		text = self._normalize_text(raw_text)
		return ParsedDocument(
			text=text,
			meta={
				"filename": path.name,
				"suffix": path.suffix.lower(),
				"parser": "docx",
			},
		)
